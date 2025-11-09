import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor
from rest_framework import status
from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from docx.enum.text import WD_ALIGN_PARAGRAPH
from apps.constructions.models import Construction



class DocumentsDetailAPIView(APIView):
    def get(self, request, pk):
        try:
            construction = Construction.objects.get(pk=pk)
        except Construction.DoesNotExist:
            return Response(
                {"error": "Construction not found"}, status=status.HTTP_404_NOT_FOUND
            )

        doc = Document()

        image_url = "https://d2bxzineatl84k.cloudfront.net/storage/files/logos/db5y8sSHEnJV4CAv3gdQdb6ZlDNZwG895zGJNuJ7.png"

        response = requests.get(image_url)
        image_bytes = BytesIO(response.content)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(12)

        # Add and center the image
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(image_bytes, width=Inches(2))
        
        heading1 = doc.add_heading(f"Especificação Técnica {construction.project_name}", level=1)
        heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading1.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
            run.font.size = Pt(15)
        
        doc.add_paragraph()  # Add space after heading1

        heading2_loc = doc.add_heading("Localização", level=2)
        for run in heading2_loc.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
        
        doc.add_paragraph()  # Add space after heading
        
        para_loc = doc.add_paragraph(construction.location)
        for run in para_loc.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()  # Add space after paragraph
        
        heading2_desc = doc.add_heading("Descrição do Empreendimento", level=2)
        for run in heading2_desc.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True
        
        doc.add_paragraph()  # Add space after heading
        
        para_desc = doc.add_paragraph(construction.description)
        for run in para_desc.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()  # Add space after paragraph

        # Especificações de materiais de acabamento por referencial
        referentials = construction.referentials.all()
        for referential in referentials:
            heading2_ref_mat = doc.add_heading(
                f"Das especificações de materiais de acabamento referente {referential.referential_name.name}",
                level=2
            )
            for run in heading2_ref_mat.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
            
            doc.add_paragraph()  # Add space after heading
            
            areas = referential.areas.all()
            for index, area in enumerate(areas):
                # Create sub-section with sequential label (a), b), c), ...)
                label = chr(ord('a') + index) + ')'
                subtitle = doc.add_paragraph(f"{label} {area.area_name.name}")
                for run in subtitle.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                doc.add_paragraph()  # Add space after subtitle
                
                # Create table with elements
                elements = area.elements.all()
                if elements.exists():
                    # Create table with header row + data rows
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Tipo de Elemento"
                    header_cells[1].text = "Material"
                    
                    # Style header cells
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Add data rows
                    for element in elements:
                        row_cells = table.add_row().cells
                        row_cells[0].text = element.element_type.name
                        # Use each material's __str__ function
                        materials_list = [str(material) for material in element.materials.all()]
                        row_cells[1].text = ", ".join(materials_list) if materials_list else "None"
                        
                        # Style data cells
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    doc.add_paragraph()  # Add space after table

        doc.add_paragraph()  # Add space before Observações section
        
        # Observações
        observations = construction.observations.all()
        if observations.exists():
            heading2_obs = doc.add_heading("Observações", level=2)
            for run in heading2_obs.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
            
            doc.add_paragraph()  # Add space after heading
            
            for obs in observations:
                para_obs = doc.add_paragraph(str(obs), style='List Bullet')
                for run in para_obs.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{construction.project_name}.docx".replace(" ", "_")

        return FileResponse(
            buffer,
            as_attachment=True,
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
